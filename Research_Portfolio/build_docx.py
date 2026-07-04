import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_cv(md_path, out_path):
    doc = docx.Document()
    
    # Read md content
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold text in list items
            text = line[2:]
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    doc.save(out_path)

if __name__ == '__main__':
    create_cv('CV_Final_US_Format.md', 'CV_Final_US_Format.docx')
